
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def parse_markdown_to_docx(md_file, docx_file):
    print(f"Reading {md_file}...")
    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    code_block = False
    
    for line in lines:
        line = line.strip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            code_block = not code_block
            continue
        
        if code_block:
            p = document.add_paragraph()
            p.style = document.styles['No Spacing']
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        if not line:
            continue

        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        
        # Images: ![Alt](Path)
        elif line.startswith('![') and '](' in line and line.endswith(')'):
            match = re.search(r'\!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                if os.path.exists(img_path):
                    print(f"Adding image: {img_path}")
                    try:
                        document.add_picture(img_path, width=Inches(6))
                    except Exception as e:
                        print(f"Error adding image {img_path}: {e}")
                        document.add_paragraph(f"[Image: {img_path}]")
                else:
                    print(f"Image not found: {img_path}")
                    document.add_paragraph(f"[Image not found: {img_path}]")
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        
        # Normal text
        else:
            p = document.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(docx_file)
    print(f"Successfully generated {docx_file}")

if __name__ == "__main__":
    md_path = '/Users/adarshainamdar/Downloads/Major_Project-main/Documents/Chapter_8_Results_and_Discussions.md'
    docx_path = '/Users/adarshainamdar/Downloads/Major_Project-main/Documents/Chapter_8_Results_and_Discussions.docx'
    parse_markdown_to_docx(md_path, docx_path)
