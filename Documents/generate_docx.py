from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx():
    document = Document()
    
    # Title
    title = document.add_heading('Project Documentation & Explanation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: Architecture ---
    document.add_heading('1. System Architecture & Design Details', level=1)
    
    document.add_heading('System Architecture Diagram', level=2)
    p = document.add_paragraph('The high-level data flow from input images to final classification.')
    
    # Diagram 1 Text Representation
    diagram_box = document.add_paragraph()
    diagram_box.paragraph_format.left_indent = Inches(0.5)
    run = diagram_box.add_run(
        "[ Input: Blood Smear Images ]\n"
        "      |\n"
        "      v\n"
        "[ Preprocessing: Resize & Normalize ]\n"
        "      |\n"
        "      v\n"
        "[ Augmentation: Flip/Rotate ]\n"
        "      |\n"
        "      v\n"
        "[ ResNet-18 Backbone CNN ] <--- (PSO Optimizes Batch Size & LR)\n"
        "      |\n"
        "      +---------------------------+\n"
        "      |                           |\n"
        "      v                           v\n"
        "[ Head 1: Cell Type ]       [ Head 2: Severity Grade ]\n"
        "      |                           |\n"
        "      v                           v\n"
        "[ ALL, AML, CLL, CML ]      [ Chronic, Accelerated, Blast ]"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    document.add_heading('Detailed Design (Class & Data Flow)', level=2)
    
    # Diagram 2 Text Representation
    diagram_box2 = document.add_paragraph()
    diagram_box2.paragraph_format.left_indent = Inches(0.5)
    run2 = diagram_box2.add_run(
        "Class: MultiTaskDataset\n"
        "  |-- __getitem__(): Returns (Image, TypeLabel, GradeLabel)\n"
        "  |\n"
        "  v\n"
        "Training Loop (Feeds Batch)\n"
        "  |\n"
        "  v\n"
        "Class: MultiTaskModel\n"
        "  |-- ResNet18 Backbone\n"
        "  |-- Forward(): Returns Logits\n"
        "  |\n"
        "  v\n"
        "Loss Function\n"
        "  |-- Total Loss = Loss_Type + 0.5 * Loss_Grade"
    )
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)

    # --- Section 2: Formulas ---
    document.add_heading('2. Formulas & Explanations', level=1)
    
    document.add_heading('Batch Size Training Data Formula (PSO)', level=2)
    document.add_paragraph('In this project, the Batch Size is optimized by Particle Swarm Optimization (PSO).')
    
    document.add_paragraph('Velocity Update:', style='List Number')
    document.add_paragraph('v(t+1) = w * v(t) + c1 * r1 * (pbest - x(t)) + c2 * r2 * (gbest - x(t))', style='Quote')
    
    document.add_paragraph('Position Update:', style='List Number')
    document.add_paragraph('x(t+1) = x(t) + v(t+1)', style='Quote')
    
    document.add_paragraph('Batch Size Mapping:', style='List Number')
    document.add_paragraph('Batch Size = Round( x_batch_dimension(t+1) )', style='Quote')
    
    document.add_heading('"Dotted Line" in Plots', level=2)
    document.add_paragraph('Solid Line: Represents Training metrics (learning from seen data).')
    document.add_paragraph('Dotted/Dashed Line: Represents Validation metrics (generalization to unseen data).')
    document.add_paragraph('Key Insight: If Train goes up but Validation stays flat/drops, it indicates Overfitting.')
    
    document.add_heading('Confusion Matrix Formula', level=2)
    document.add_paragraph('Accuracy = (TP + TN) / (TP + TN + FP + FN)')
    p = document.add_paragraph()
    p.add_run('TP: True Positive, TN: True Negative, FP: False Positive, FN: False Negative').italic = True

    # --- Section 3: Requirements ---
    document.add_heading('3. Software Requirements', level=1)
    reqs = [
        "OS: macOS (Sonoma/Sequoia) with Apple Silicon (M1/M2/M3)",
        "Python: 3.10+",
        "Dependencies: torch, torchvision, pyswarm, opencv-python, scikit-learn, matplotlib, seaborn, tqdm"
    ]
    for r in reqs:
        document.add_paragraph(r, style='List Bullet')

    # --- Section 4: Chapter 8 ---
    document.add_heading('4. Chapter 8: Results and Discussions', level=1)
    
    document.add_heading('8.1 Experimental Setup', level=2)
    document.add_paragraph('Data split: 70% Train, 15% Val, 15% Test. Classes: ALL, AML, CLL, CML.')
    
    document.add_heading('8.2 Baseline Results', level=2)
    document.add_paragraph('Model: ResNet-18 (Standard).')
    document.add_paragraph('Peak Validation Accuracy: 93.23% (Epoch 62).')
    document.add_paragraph('Observation: Showed fluctuations suggesting learning rate sensitivity.')
    
    document.add_heading('8.3 Multi-Task PSO Results', level=2)
    document.add_paragraph('Model: Multi-Task + PSO Optimization.')
    document.add_paragraph('PSO Optimal Hyperparameters: Batch Size = 32, LR = 0.000163.')
    document.add_paragraph('Peak Validation Accuracy: 96.88% (Epoch 34).')
    document.add_paragraph('Improvement: Higher accuracy (+3.65%) and faster convergence compared to baseline.')
    
    document.add_heading('8.4 Discussion', level=2)
    document.add_paragraph('The improvement is attributed to the regularization effect of Multi-Task Learning (Auxiliary Grade task) and the precise hyperparameter tuning by PSO. The model effectively distinguishes Acute vs Chronic types.')

    # Save
    out_path = '/Users/adarshainamdar/Downloads/Major_Project-main/Documents/Project_Documentation.docx'
    document.save(out_path)
    print(f"Document saved to {out_path}")

if __name__ == "__main__":
    create_docx()
