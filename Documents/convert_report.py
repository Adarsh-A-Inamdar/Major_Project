import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_docx(md_file, docx_file):
    document = Document()
    
    # Set style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
         
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Remove the number and dot
            text = re.sub(r'^\d+\.\s*', '', line)
            document.add_paragraph(text, style='List Number')
            
        # Code blocks (simple handling)
        elif line.startswith('```'):
            continue # Skip code fences for now, or maybe handle them
        
        # Normal text
        else:
            p = document.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(docx_file)
    print(f"Successfully generated {docx_file}")

if __name__ == "__main__":
    parse_markdown_to_docx(
        '/Users/adarshainamdar/.gemini/antigravity/brain/b38fa485-155c-4e15-b19b-406d81684806/project_report.md',
        '/Users/adarshainamdar/Downloads/Major_Project-main/Major_Project_Report.docx'
    )
